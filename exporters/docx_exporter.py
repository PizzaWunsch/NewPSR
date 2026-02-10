# exporters/docx_exporter.py
from __future__ import annotations

import os
import json
from datetime import datetime
from typing import Any, Dict

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export_docx(out_dir: str, title: str = "Schritt-für-Schritt Anleitung"):
    steps_path = os.path.join(out_dir, "steps.json")
    if not os.path.exists(steps_path):
        raise FileNotFoundError(f"steps.json not found in {out_dir}")

    with open(steps_path, "r", encoding="utf-8") as f:
        data: Dict[str, Any] = json.load(f)

    events = data.get("events", [])
    monitors = {m["index"]: m for m in data.get("monitors", [])}
    delay_ms = int(data.get("screenshot_delay_ms") or 0)

    def mon_label(idx):
        if not idx or idx not in monitors:
            return None
        m = monitors[idx]
        return f"Monitor {idx} ({m['width']}×{m['height']})"

    doc = Document()

    t = doc.add_paragraph(title)
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(20)
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT

    meta = doc.add_paragraph(f"Erstellt: {datetime.now().strftime('%d.%m.%Y %H:%M')}  ·  Screenshot-Verzögerung: {delay_ms} ms")
    meta.runs[0].font.size = Pt(10)

    doc.add_paragraph("")

    step_no = 0
    for e in events:
        if e.get("kind") in ("mouse_click", "key_press", "text_input"):
            text = (e.get("instruction") or e.get("detail") or "").strip()
            if not text:
                continue
            step_no += 1

            p = doc.add_paragraph()
            r = p.add_run(f"{step_no}. {text}")
            r.bold = True
            r.font.size = Pt(12)

            ml = mon_label(e.get("monitor_index"))
            if ml:
                pm = doc.add_paragraph(ml)
                pm.runs[0].font.size = Pt(10)

            ss = e.get("screenshot")
            if ss:
                abs_img = os.path.join(out_dir, ss)
                if os.path.exists(abs_img):
                    doc.add_picture(abs_img, width=Inches(6.5))

            doc.add_paragraph("")

    out_path = os.path.join(out_dir, "anleitung.docx")
    doc.save(out_path)
    return out_path